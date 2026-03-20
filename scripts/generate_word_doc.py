#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成系统需求规格说明书Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_with_style(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_style(doc, text, bold=False, size=11):
    """添加段落"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
    return p

def create_srs_document():
    """创建系统需求规格说明书"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    
    # 标题页
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('天津华能票据审核系统')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '宋体'
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('系统需求规格说明书')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = '宋体'
    
    doc.add_paragraph()
    
    # 文档信息表
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    cells = info_table.rows[0].cells
    cells[0].text = '文档版本'
    cells[1].text = 'v1.0'
    
    cells = info_table.rows[1].cells
    cells[0].text = '更新日期'
    cells[1].text = '2025年2月28日'
    
    cells = info_table.rows[2].cells
    cells[0].text = '文档状态'
    cells[1].text = '正式版'
    
    cells = info_table.rows[3].cells
    cells[0].text = '编制部门'
    cells[1].text = '系统分析团队'
    
    cells = info_table.rows[4].cells
    cells[0].text = '适用范围'
    cells[1].text = '开发、测试、部署和运维人员'
    
    doc.add_page_break()
    
    # 目录
    add_heading_with_style(doc, '目录', 1)
    toc_items = [
        '1. 文档概述',
        '2. 系统概述',
        '3. 用户角色与权限',
        '4. 功能需求',
        '5. 非功能需求',
        '6. 系统架构',
        '7. API接口规范',
        '8. 部署与运维',
        '9. 测试需求',
        '10. 版本历史',
        '11. 附录'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. 文档概述
    add_heading_with_style(doc, '1. 文档概述', 1)
    
    add_heading_with_style(doc, '1.1 文档目的', 2)
    add_paragraph_with_style(doc, '本文档定义了天津华能票据审核系统的功能需求、非功能需求、系统架构和技术规范，为系统开发、测试和维护提供指导。')
    
    add_heading_with_style(doc, '1.2 适用范围', 2)
    add_paragraph_with_style(doc, '本规格说明书适用于天津华能票据审核系统的所有开发、测试、部署和运维人员。')
    
    add_heading_with_style(doc, '1.3 文档版本', 2)
    version_table = doc.add_table(rows=4, cols=2)
    version_table.style = 'Light Grid Accent 1'
    
    cells = version_table.rows[0].cells
    cells[0].text = '版本号'
    cells[1].text = 'v1.0'
    
    cells = version_table.rows[1].cells
    cells[0].text = '更新日期'
    cells[1].text = '2025年2月28日'
    
    cells = version_table.rows[2].cells
    cells[0].text = '状态'
    cells[1].text = '正式版'
    
    cells = version_table.rows[3].cells
    cells[0].text = '备注'
    cells[1].text = '初始版本'
    
    doc.add_page_break()
    
    return doc

# 继续在下一部分添加内容
if __name__ == '__main__':
    doc = create_srs_document()
    doc.save('docs/系统需求规格说明书.docx')
    print('Word文档已生成：docs/系统需求规格说明书.docx')


def add_system_overview(doc):
    """添加系统概述部分"""
    add_heading_with_style(doc, '2. 系统概述', 1)
    
    add_heading_with_style(doc, '2.1 系统定义', 2)
    add_paragraph_with_style(doc, '天津华能票据审核系统是一套企业级的票据审批和项目管理平台，支持多级审批流程、项目材料管理、采购合同管理等功能。系统采用移动端优先设计，支持在手机、平板和电脑上使用。')
    
    add_heading_with_style(doc, '2.2 系统目标', 2)
    objectives = [
        '规范企业票据审批流程，提高审批效率',
        '实现多级审批管理，确保财务合规',
        '支持项目和材料清单管理',
        '提供完整的审批历史和数据追溯',
        '支持移动端访问，提升用户体验'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    add_heading_with_style(doc, '2.3 主要功能模块', 2)
    modules = [
        '用户认证与权限管理 - 登录、密码管理、角色权限控制',
        '项目管理 - 项目创建、编辑、查看、删除',
        '材料清单管理 - 材料、施工清单、调试清单的管理',
        '采购合同管理 - 供应商合同的创建和管理',
        '票据管理 - 票据提交、修改、查看',
        '审批流程 - 多级审批、拒绝、重提',
        '归档管理 - 已完成票据的归档',
        '系统管理 - 用户管理、审批流程配置、事由分类管理'
    ]
    for module in modules:
        doc.add_paragraph(module, style='List Number')
    
    doc.add_page_break()

def add_roles_and_permissions(doc):
    """添加用户角色与权限部分"""
    add_heading_with_style(doc, '3. 用户角色与权限', 1)
    
    add_heading_with_style(doc, '3.1 角色定义', 2)
    
    # 角色定义表
    role_table = doc.add_table(rows=14, cols=4)
    role_table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = role_table.rows[0].cells
    header_cells[0].text = '角色ID'
    header_cells[1].text = '角色名称'
    header_cells[2].text = '权限范围'
    header_cells[3].text = '主要功能'
    
    # 设置表头背景色
    for cell in header_cells:
        set_cell_background(cell, 'D3D3D3')
    
    roles_data = [
        ('admin', '管理员', '系统级', '用户管理、审批流程配置、事由分类管理、公司名称设置'),
        ('chairman', '董事长', '企业级', '项目看板、高级审批、项目审批'),
        ('vice_chairman', '副董事长', '企业级', '项目看板、后勤类审批'),
        ('gm', '总经理', '企业级', '项目看板、工程类审批'),
        ('procurement_manager', '采购经理', '项目级', '项目管理、材料清单、采购合同'),
        ('cost_manager', '造价经理', '项目级', '项目管理、材料清单、采购合同'),
        ('project_manager', '项目经理', '项目级', '项目查看、材料清单管理'),
        ('approver1', '一级审核', '票据级', '票据一级审批'),
        ('approver2', '二级审核', '票据级', '票据二级审批'),
        ('approver3', '三级审核', '票据级', '票据三级审批'),
        ('finance_manager', '财务经理', '票据级', '票据审批、票据归档'),
        ('accountant', '会计', '票据级', '票据归档、财务审核'),
        ('staff', '普通员工', '个人级', '提交票据、查看自己的票据')
    ]
    
    for i, (role_id, role_name, scope, functions) in enumerate(roles_data, 1):
        cells = role_table.rows[i].cells
        cells[0].text = role_id
        cells[1].text = role_name
        cells[2].text = scope
        cells[3].text = functions
    
    doc.add_page_break()

def add_functional_requirements(doc):
    """添加功能需求部分"""
    add_heading_with_style(doc, '4. 功能需求', 1)
    
    add_heading_with_style(doc, '4.1 用户认证与权限管理', 2)
    
    add_heading_with_style(doc, '4.1.1 登录功能', 3)
    add_paragraph_with_style(doc, '需求：用户通过用户ID和密码登录系统')
    add_paragraph_with_style(doc, '输入：用户ID、密码')
    add_paragraph_with_style(doc, '输出：JWT令牌、用户信息（ID、姓名、角色）')
    add_paragraph_with_style(doc, '验证：')
    doc.add_paragraph('用户ID和密码必填', style='List Bullet')
    doc.add_paragraph('密码验证失败返回错误提示', style='List Bullet')
    doc.add_paragraph('登录成功返回7天有效期的JWT令牌', style='List Bullet')
    
    add_heading_with_style(doc, '4.1.2 密码管理', 3)
    add_paragraph_with_style(doc, '修改密码：用户可修改自己的密码，需验证原密码')
    add_paragraph_with_style(doc, '重置密码：管理员可重置任意用户密码为默认值 123456')
    add_paragraph_with_style(doc, '默认密码：')
    doc.add_paragraph('管理员：admin123', style='List Bullet')
    doc.add_paragraph('其他用户：123456', style='List Bullet')
    
    add_heading_with_style(doc, '4.1.3 权限控制', 3)
    add_paragraph_with_style(doc, '基于角色的访问控制（RBAC）：根据用户角色限制功能访问')
    add_paragraph_with_style(doc, 'API级权限：所有需要权限的API需验证JWT令牌和用户角色')
    add_paragraph_with_style(doc, '页面级权限：前端根据用户角色显示/隐藏相应功能')
    
    add_heading_with_style(doc, '4.2 项目管理', 2)
    
    add_heading_with_style(doc, '4.2.1 项目创建', 3)
    add_paragraph_with_style(doc, '权限：采购经理、造价经理、项目经理')
    add_paragraph_with_style(doc, '必填字段：')
    doc.add_paragraph('项目名称', style='List Bullet')
    doc.add_paragraph('甲方单位', style='List Bullet')
    doc.add_paragraph('合同金额', style='List Bullet')
    
    add_heading_with_style(doc, '4.2.2 项目编辑', 3)
    add_paragraph_with_style(doc, '权限：项目创建者、采购经理、造价经理')
    add_paragraph_with_style(doc, '功能：修改项目信息、上传/删除附件')
    
    add_heading_with_style(doc, '4.3 材料清单管理', 2)
    
    add_heading_with_style(doc, '4.3.1 材料添加', 3)
    add_paragraph_with_style(doc, '权限：采购经理、造价经理、项目经理')
    add_paragraph_with_style(doc, '必填字段：')
    doc.add_paragraph('材料名称', style='List Bullet')
    doc.add_paragraph('数量', style='List Bullet')
    doc.add_paragraph('单价', style='List Bullet')
    add_paragraph_with_style(doc, '自动计算：总价 = 数量 × 单价')
    add_paragraph_with_style(doc, '自动完成：输入材料名称时显示建议，点击建议自动填充规格、单位、单价')
    
    add_heading_with_style(doc, '4.4 采购合同管理', 2)
    
    add_heading_with_style(doc, '4.4.1 合同创建', 3)
    add_paragraph_with_style(doc, '权限：采购经理、造价经理')
    add_paragraph_with_style(doc, '必填字段：')
    doc.add_paragraph('合同名称', style='List Bullet')
    doc.add_paragraph('供应商名称', style='List Bullet')
    
    add_heading_with_style(doc, '4.5 票据管理', 2)
    
    add_heading_with_style(doc, '4.5.1 票据提交', 3)
    add_paragraph_with_style(doc, '权限：普通员工')
    add_paragraph_with_style(doc, '必填字段：')
    doc.add_paragraph('票据标题', style='List Bullet')
    doc.add_paragraph('金额', style='List Bullet')
    doc.add_paragraph('事由分类（一级分类 > 二级项目）', style='List Bullet')
    doc.add_paragraph('日期', style='List Bullet')
    add_paragraph_with_style(doc, '支持格式：JPG、JPEG、PNG、WEBP')
    add_paragraph_with_style(doc, '单张图片最大：10MB，每张票据最多上传5张图片')
    
    add_heading_with_style(doc, '4.6 审批流程', 2)
    
    add_heading_with_style(doc, '4.6.1 审批路径', 3)
    add_paragraph_with_style(doc, '工程款项（事由分类包含"工程"）：')
    doc.add_paragraph('总经理 → 董事长', style='List Bullet')
    add_paragraph_with_style(doc, '后勤花费（事由分类包含"后勤"）：')
    doc.add_paragraph('副董事长 → 董事长', style='List Bullet')
    add_paragraph_with_style(doc, '其他类票据：')
    doc.add_paragraph('一级审核 → 二级审核 → 三级审核 → 董事长', style='List Bullet')
    
    add_heading_with_style(doc, '4.6.2 审批操作', 3)
    add_paragraph_with_style(doc, '通过：票据进入下一审批环节')
    add_paragraph_with_style(doc, '拒绝：填写拒绝原因，票据状态变为"已拒绝"或退回上一级')
    add_paragraph_with_style(doc, '权限：只有当前步骤的审批人或管理员可操作')
    
    doc.add_page_break()

def add_non_functional_requirements(doc):
    """添加非功能需求部分"""
    add_heading_with_style(doc, '5. 非功能需求', 1)
    
    add_heading_with_style(doc, '5.1 性能需求', 2)
    add_paragraph_with_style(doc, 'API响应时间 < 2秒')
    add_paragraph_with_style(doc, '支持至少100个并发用户')
    add_paragraph_with_style(doc, 'SQLite数据库，支持至少10万条票据记录')
    
    add_heading_with_style(doc, '5.2 安全需求', 2)
    add_paragraph_with_style(doc, 'JWT令牌认证，7天有效期')
    add_paragraph_with_style(doc, '密码以明文存储（开发版本），生产环境应加密')
    add_paragraph_with_style(doc, '支持跨域请求配置')
    add_paragraph_with_style(doc, '所有修改操作需验证用户权限')
    
    add_heading_with_style(doc, '5.3 可用性需求', 2)
    add_paragraph_with_style(doc, '系统可用性：99%以上')
    add_paragraph_with_style(doc, '定期备份数据库')
    add_paragraph_with_style(doc, '友好的错误提示')
    
    add_heading_with_style(doc, '5.4 兼容性需求', 2)
    add_paragraph_with_style(doc, '支持Chrome、Safari、Firefox等现代浏览器')
    add_paragraph_with_style(doc, '支持iOS、Android系统')
    add_paragraph_with_style(doc, '支持微信内置浏览器访问')
    
    doc.add_page_break()

def add_system_architecture(doc):
    """添加系统架构部分"""
    add_heading_with_style(doc, '6. 系统架构', 1)
    
    add_heading_with_style(doc, '6.1 技术栈', 2)
    add_paragraph_with_style(doc, '前端：React 19 + Vite + Tailwind CSS + Material-UI')
    add_paragraph_with_style(doc, '后端：Node.js + Express')
    add_paragraph_with_style(doc, '数据库：SQLite')
    add_paragraph_with_style(doc, '认证：JWT')
    add_paragraph_with_style(doc, '文件上传：Multer')
    
    add_heading_with_style(doc, '6.2 主要数据库表', 2)
    
    add_heading_with_style(doc, '6.2.1 users 表', 3)
    add_paragraph_with_style(doc, '用户信息表，存储用户ID、姓名、角色、密码')
    
    add_heading_with_style(doc, '6.2.2 bills 表', 3)
    add_paragraph_with_style(doc, '票据表，存储票据信息、审批流程、审批历史')
    
    add_heading_with_style(doc, '6.2.3 projects 表', 3)
    add_paragraph_with_style(doc, '项目表，存储项目基本信息、预算、状态等')
    
    add_heading_with_style(doc, '6.2.4 materials 表', 3)
    add_paragraph_with_style(doc, '材料清单表，存储项目材料信息')
    
    add_heading_with_style(doc, '6.2.5 supplier_contracts 表', 3)
    add_paragraph_with_style(doc, '采购合同表，存储供应商合同信息')
    
    add_heading_with_style(doc, '6.2.6 settings 表', 3)
    add_paragraph_with_style(doc, '系统设置表，存储公司名称、审批阈值等配置')
    
    doc.add_page_break()

def add_api_specification(doc):
    """添加API接口规范部分"""
    add_heading_with_style(doc, '7. API接口规范', 1)
    
    add_heading_with_style(doc, '7.1 认证接口', 2)
    
    add_heading_with_style(doc, '7.1.1 POST /api/login', 3)
    add_paragraph_with_style(doc, '登录接口')
    add_paragraph_with_style(doc, '请求：{ id: string, password: string }')
    add_paragraph_with_style(doc, '响应：{ id, name, role, token }')
    
    add_heading_with_style(doc, '7.1.2 POST /api/user/change-password', 3)
    add_paragraph_with_style(doc, '修改密码')
    add_paragraph_with_style(doc, '权限：本人或管理员')
    add_paragraph_with_style(doc, '请求：{ id, oldPassword, newPassword }')
    add_paragraph_with_style(doc, '响应：{ ok: true }')
    
    add_heading_with_style(doc, '7.2 用户接口', 2)
    
    add_heading_with_style(doc, '7.2.1 GET /api/users', 3)
    add_paragraph_with_style(doc, '获取所有用户')
    add_paragraph_with_style(doc, '响应：[{ id, name, role }, ...]')
    
    add_heading_with_style(doc, '7.2.2 POST /api/users', 3)
    add_paragraph_with_style(doc, '添加用户')
    add_paragraph_with_style(doc, '权限：管理员')
    add_paragraph_with_style(doc, '请求：{ id, name, role, password }')
    add_paragraph_with_style(doc, '响应：{ success: true, message: string }')
    
    add_heading_with_style(doc, '7.3 项目接口', 2)
    
    add_heading_with_style(doc, '7.3.1 GET /api/projects', 3)
    add_paragraph_with_style(doc, '获取所有项目')
    add_paragraph_with_style(doc, '响应：[{ id, code, name, client, totalBudget, ... }, ...]')
    
    add_heading_with_style(doc, '7.3.2 POST /api/projects', 3)
    add_paragraph_with_style(doc, '创建项目')
    add_paragraph_with_style(doc, '权限：采购经理、造价经理、项目经理')
    add_paragraph_with_style(doc, '请求：{ name, client, contractNo, totalBudget, ... }')
    add_paragraph_with_style(doc, '响应：{ id, code, name, ... }')
    
    add_heading_with_style(doc, '7.4 票据接口', 2)
    
    add_heading_with_style(doc, '7.4.1 GET /api/bills', 3)
    add_paragraph_with_style(doc, '获取所有票据')
    add_paragraph_with_style(doc, '响应：[{ id, title, amount, category, status, ... }, ...]')
    
    add_heading_with_style(doc, '7.4.2 POST /api/bill', 3)
    add_paragraph_with_style(doc, '创建票据')
    add_paragraph_with_style(doc, '权限：普通员工')
    add_paragraph_with_style(doc, '请求：{ title, amount, category, date, projectId?, images? }')
    add_paragraph_with_style(doc, '响应：{ id, title, amount, ... }')
    
    add_heading_with_style(doc, '7.4.3 POST /api/bill/approve', 3)
    add_paragraph_with_style(doc, '审批票据')
    add_paragraph_with_style(doc, '权限：审批人')
    add_paragraph_with_style(doc, '请求：{ id, action: "approve"|"reject", reason? }')
    add_paragraph_with_style(doc, '响应：{ ok: true }')
    
    doc.add_page_break()

def add_deployment_and_maintenance(doc):
    """添加部署与运维部分"""
    add_heading_with_style(doc, '8. 部署与运维', 1)
    
    add_heading_with_style(doc, '8.1 部署方式', 2)
    
    add_heading_with_style(doc, '8.1.1 开发环境', 3)
    add_paragraph_with_style(doc, 'npm install')
    add_paragraph_with_style(doc, 'npm run dev          # 前端开发服务器')
    add_paragraph_with_style(doc, 'npm run server       # 后端服务器')
    
    add_heading_with_style(doc, '8.1.2 生产环境', 3)
    add_paragraph_with_style(doc, 'npm install')
    add_paragraph_with_style(doc, 'npm run build        # 构建前端')
    add_paragraph_with_style(doc, 'npm run server       # 启动后端服务器')
    
    add_heading_with_style(doc, '8.2 环境变量', 2)
    add_paragraph_with_style(doc, 'JWT_SECRET：JWT签名密钥（默认：dev-secret）')
    add_paragraph_with_style(doc, 'ALLOW_ORIGINS：允许的CORS源（逗号分隔）')
    add_paragraph_with_style(doc, 'ALLOW_DEV_RESET：是否启用开发重置接口（默认：1）')
    
    add_heading_with_style(doc, '8.3 数据库备份', 2)
    add_paragraph_with_style(doc, '数据库位置：server/data/app.db')
    add_paragraph_with_style(doc, '备份文件：server/data/app.db.backup')
    add_paragraph_with_style(doc, '建议定期备份数据库文件')
    
    add_heading_with_style(doc, '8.4 文件上传', 2)
    add_paragraph_with_style(doc, '上传目录：server/data/uploads')
    add_paragraph_with_style(doc, '支持的文件类型：图片（JPG、PNG、WEBP等）')
    add_paragraph_with_style(doc, '最大文件大小：10MB')
    
    doc.add_page_break()

def add_testing_requirements(doc):
    """添加测试需求部分"""
    add_heading_with_style(doc, '9. 测试需求', 1)
    
    add_heading_with_style(doc, '9.1 功能测试', 2)
    doc.add_paragraph('用户登录、密码修改', style='List Bullet')
    doc.add_paragraph('项目创建、编辑、删除', style='List Bullet')
    doc.add_paragraph('材料清单管理', style='List Bullet')
    doc.add_paragraph('采购合同管理', style='List Bullet')
    doc.add_paragraph('票据提交、审批、拒绝、重提', style='List Bullet')
    doc.add_paragraph('归档操作', style='List Bullet')
    
    add_heading_with_style(doc, '9.2 权限测试', 2)
    doc.add_paragraph('验证各角色的功能访问权限', style='List Bullet')
    doc.add_paragraph('验证API权限控制', style='List Bullet')
    
    add_heading_with_style(doc, '9.3 性能测试', 2)
    doc.add_paragraph('API响应时间', style='List Bullet')
    doc.add_paragraph('并发用户处理能力', style='List Bullet')
    doc.add_paragraph('大数据量处理', style='List Bullet')
    
    add_heading_with_style(doc, '9.4 安全测试', 2)
    doc.add_paragraph('SQL注入防护', style='List Bullet')
    doc.add_paragraph('XSS防护', style='List Bullet')
    doc.add_paragraph('CSRF防护', style='List Bullet')
    doc.add_paragraph('密码安全', style='List Bullet')
    
    doc.add_page_break()

def add_version_history(doc):
    """添加版本历史部分"""
    add_heading_with_style(doc, '10. 版本历史', 1)
    
    version_table = doc.add_table(rows=4, cols=3)
    version_table.style = 'Light Grid Accent 1'
    
    header_cells = version_table.rows[0].cells
    header_cells[0].text = '版本'
    header_cells[1].text = '日期'
    header_cells[2].text = '主要变更'
    
    for cell in header_cells:
        set_cell_background(cell, 'D3D3D3')
    
    cells = version_table.rows[1].cells
    cells[0].text = 'v0.1'
    cells[1].text = '2025-01-16'
    cells[2].text = '初始版本'
    
    cells = version_table.rows[2].cells
    cells[0].text = 'v0.2'
    cells[1].text = '2025-01-20'
    cells[2].text = '添加造价经理角色、材料自动完成功能'
    
    cells = version_table.rows[3].cells
    cells[0].text = 'v1.0'
    cells[1].text = '2025-02-28'
    cells[2].text = '完整的系统需求规格说明书'
    
    doc.add_page_break()

def add_appendix(doc):
    """添加附录部分"""
    add_heading_with_style(doc, '11. 附录', 1)
    
    add_heading_with_style(doc, '11.1 术语表', 2)
    
    terms_table = doc.add_table(rows=6, cols=2)
    terms_table.style = 'Light Grid Accent 1'
    
    header_cells = terms_table.rows[0].cells
    header_cells[0].text = '术语'
    header_cells[1].text = '定义'
    
    for cell in header_cells:
        set_cell_background(cell, 'D3D3D3')
    
    terms = [
        ('JWT', 'JSON Web Token，用于身份认证'),
        ('RBAC', 'Role-Based Access Control，基于角色的访问控制'),
        ('CORS', 'Cross-Origin Resource Sharing，跨域资源共享'),
        ('SQLite', '轻量级关系型数据库'),
        ('Multer', 'Node.js文件上传中间件')
    ]
    
    for i, (term, definition) in enumerate(terms, 1):
        cells = terms_table.rows[i].cells
        cells[0].text = term
        cells[1].text = definition
    
    add_heading_with_style(doc, '11.2 参考文档', 2)
    doc.add_paragraph('用户手册：docs/用户手册.md', style='List Bullet')
    doc.add_paragraph('项目README：README.md', style='List Bullet')
    doc.add_paragraph('API文档：待补充', style='List Bullet')
    
    # 添加页脚
    doc.add_paragraph()
    footer = doc.add_paragraph('本文档由系统分析团队编制，最后更新于2025年2月28日。')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(10)
        run.font.italic = True

def main():
    """主函数"""
    doc = create_srs_document()
    
    # 添加各部分内容
    add_system_overview(doc)
    add_roles_and_permissions(doc)
    add_functional_requirements(doc)
    add_non_functional_requirements(doc)
    add_system_architecture(doc)
    add_api_specification(doc)
    add_deployment_and_maintenance(doc)
    add_testing_requirements(doc)
    add_version_history(doc)
    add_appendix(doc)
    
    # 保存文档
    doc.save('docs/系统需求规格说明书.docx')
    print('✓ Word文档已生成：docs/系统需求规格说明书.docx')

if __name__ == '__main__':
    main()
